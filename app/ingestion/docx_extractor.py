"""Word (.docx) extraction: paragraphs + tables."""
from __future__ import annotations

from pathlib import Path

from app.core.logging_config import get_logger

log = get_logger(__name__)


def extract_docx(path: Path, doc) -> None:
    try:
        from docx import Document
    except Exception as exc:  # noqa: BLE001
        log.warning("python-docx not available: %s", exc)
        return

    try:
        document = Document(str(path))
    except Exception as exc:  # noqa: BLE001
        log.warning("Failed to open Word doc %s: %s", path.name, exc)
        return

    parts: list[str] = []
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in document.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows:
            doc.tables.append(rows)
            for r in rows:
                parts.append("\t".join(r))

    doc.pages = ["\n".join(parts)]
    doc.full_text = "\n".join(parts).strip()
